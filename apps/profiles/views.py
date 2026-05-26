import json
import io
import base64
from rest_framework.request import Request
from rest_framework.response import Response
from drf_spectacular.utils import extend_schema
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser
from rest_framework.permissions import AllowAny
from utils.user_dto_view import UserDTOView
from apps.accounts.permissions import RolePermission
from apps.accounts.authentication import CustomJWTAuthentication
from .services import update_user_profile, calculate_profile_completion
from .services import enrich_profile_data
from apps.accounts.models import User
import pdfminer.high_level
from .services import is_profile_complete
from openai import AzureOpenAI
from django.conf import settings

import docx
import fitz # PyMuPDF
import easyocr
import numpy as np
import cv2
from PIL import Image as PILImage

def extract_docx(file):
    doc = docx.Document(file)
    full_text = []
    
    # Extract from headers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

    # Extract from main body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                                   
    # Extract from footers
    for section in doc.sections:
        if section.footer:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

    return "\n".join(full_text)

client = AzureOpenAI(
    api_key=settings.AZURE_OPENAI_API_KEY,
    azure_endpoint=settings.AZURE_OPENAI_ENDPOINT,
    api_version="2024-02-15-preview"
)

# Initialize EasyOCR reader once at startup (or lazily) to avoid per-request latency
_EASYOCR_READER = None

def get_easyocr_reader():
    global _EASYOCR_READER
    if _EASYOCR_READER is None:
        try:
            import easyocr
            # Note: gpu=False for better compatibility in CPU environments
            _EASYOCR_READER = easyocr.Reader(['en'], gpu=False, verbose=False)
        except Exception as e:
            print(f"Failed to initialize EasyOCR: {e}")
    return _EASYOCR_READER

class GetProfileView(UserDTOView):
    """
    API to retrieve user profile JSON blob.
    """

    @extend_schema(request=None, responses={200: "Profile JSON with completion metadata."})
    def get(self, request: Request) -> Response:
        target_user_id = self.user_dto.id
        
        success, profile_or_reason = update_user_profile(
            target_user_id, None, retrieve=True
        )

        if success:
            profile_data = profile_or_reason or {}

            # Enrich profile with authoritative User-model fields
            profile_data = enrich_profile_data(target_user_id, profile_data)

            completion_percentage, missing_sections = calculate_profile_completion(profile_data)
            print(f"[Profile GET] Retrieved profile of user {target_user_id}")
            print(f"[Profile GET] Completion: {completion_percentage}%, Missing sections: {missing_sections}")
            return Response({
                "profile": profile_data,
                "completion_percentage": completion_percentage,
                "is_complete": is_profile_complete(profile_data),
                "missing_sections": missing_sections,
            }, status=200)

        print(f"[Profile GET] Error: {profile_or_reason}")
        return Response({"error": profile_or_reason}, status=404)


from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt

@method_decorator(csrf_exempt, name='dispatch')
class UpdateProfileView(UserDTOView):
    """
    API to update user profile JSON blob.
    """

    @extend_schema(
        request=None, responses={200: "Profile updated successfully."}
    )
    def post(self, request: Request) -> Response:
        if not isinstance(request.data, dict):
            return Response({"error": "Payload must be a JSON object."}, status=400)

        target_user_id = self.user_dto.id
        print(f"[Profile POST] Updating for user {target_user_id}. Data keys: {list(request.data.keys())}")

        # The frontend usually sends { "profile": { ...data... } }
        # Sometimes it can be double nested due to context updates.
        # Recursively unwrap until we get to the core data.
        new_data = request.data
        while isinstance(new_data, dict) and "profile" in new_data and len(new_data) == 1:
            new_data = new_data["profile"]
            
        if not isinstance(new_data, dict):
            return Response({"error": "Profile data must be a JSON object."}, status=400)

        # Sync firstName/lastName from profile personalDetails to User model
        try:
            personal = new_data.get("personalDetails", {})
            if isinstance(personal, dict):
                first_name = personal.get("firstName", "").strip()
                last_name = personal.get("lastName", "").strip()
                if first_name or last_name:
                    User.objects.filter(id=target_user_id).update(
                        **{k: v for k, v in {"first_name": first_name, "last_name": last_name}.items() if v}
                    )
        except Exception as e:
            print(f"[Profile POST] Failed to sync name: {e}")

        # Sync academicLevel from educational to User.academic_level and strip from profile blob
        import copy
        processed_profile = copy.deepcopy(new_data)
        try:
            educational = processed_profile.get("educational", {})
            if isinstance(educational, dict) and "academicLevel" in educational:
                academic_level_label = educational.pop("academicLevel")
                # Map frontend display labels to backend model choices
                label_to_value = {label: val for val, label in User.AcademicLevel.choices}
                academic_level_val = label_to_value.get(academic_level_label, academic_level_label)
                User.objects.filter(id=target_user_id).update(academic_level=academic_level_val)
                print(f"[Profile POST] Synced academicLevel: {academic_level_val}")
        except Exception as e:
            print(f"[Profile POST] Failed to sync academicLevel: {e}")

        # Always save in the wrapped format { "profile": { ... } } for consistency
        data_to_save = {"profile": processed_profile}
        
        success, reason = update_user_profile(target_user_id, data_to_save)

        if success:
            return Response({"message": "Profile updated successfully."}, status=200)

        return Response({"error": str(reason)}, status=400)


class ResumeParserView(UserDTOView):

    parser_classes = [MultiPartParser]

    def post(self, request):
        # FIX 1: Retrieve uploaded file from request.FILES
        uploaded_file = request.FILES.get("file")

        if not uploaded_file:
            return Response({"error": "No file provided."}, status=400)

        # FIX 2: Moved file-type handling inside the method with correct indentation
        file_name = uploaded_file.name.lower()

        if file_name.endswith(".pdf"):
            text = pdfminer.high_level.extract_text(uploaded_file.file)

        elif file_name.endswith(".docx"):
            text = extract_docx(uploaded_file)

        else:
            return Response({"error": "Unsupported file type"}, status=400)

        # FIX 3: Added the missing client.chat.completions.create(...) call and assignment
        response = client.chat.completions.create(
            model=settings.AZURE_OPENAI_DEPLOYMENT,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": "Extract structured student profile data from resumes."
                },
                {
                    "role": "user",
                    "content": f"""
                    Extract structured profile data from this resume text:         
                    {text}
                    Return JSON with fields:
                        first_name
                        last_name
                        email
                        phone
                        gender
                        dob
                        city
                        zip_code
                        citizenship
                        address
                        mother_profession
                        father_profession
                        institution
                        degree
                        major
                        cgpa
                        start_year
                        end_year
                        current_year
                        skills
                        projects
                        experience
                        activities
                        certifications
                    """  # FIX 4: Closed the f-string properly
                }
            ],
            temperature=0
        )
        import json
        try:
            parsed = json.loads(response.choices[0].message.content)
        except:
            parsed = {}

        return Response({
            "personal": parsed
        })

class TranscriptParserView(UserDTOView):
    parser_classes = [MultiPartParser]
    allow_public = False # Force authentication

    def dispatch(self, request, *args, **kwargs):
        print(f"[TranscriptParser] Dispatching {request.method} {request.path}")
        return super().dispatch(request, *args, **kwargs)

    def post(self, request):
        print(f"[TranscriptParser] Received request. Files: {request.FILES.keys()}")
        uploaded_file = request.FILES.get("file")

        if not uploaded_file:
            return Response({"error": "No file provided."}, status=400)

        file_name = uploaded_file.name.lower()
        text = ""
        images_base64 = [] # For Vision API

        try:
            file_content = uploaded_file.read()
            if file_name.endswith(".pdf"):
                # Extract text for maximum speed and accuracy
                print(f"Processing PDF for text extraction...")
                doc = fitz.open(stream=file_content, filetype="pdf")
                pdf_text_parts = []
                for page_num in range(min(len(doc), 10)):
                    page = doc.load_page(page_num)
                    pdf_text_parts.append(page.get_text())
                
                text = "\n".join(pdf_text_parts)
                
                # Only use images / Vision OCR if the PDF has very little selectable text (e.g., scanned PDF)
                if len(text.strip()) < 200:
                    print(f"PDF contains minimal selectable text ({len(text)} chars). Extracting high-res page images for Vision OCR...")
                    for page_num in range(min(len(doc), 10)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                        img_bytes = pix.tobytes("jpg", jpg_quality=85)
                        images_base64.append(base64.b64encode(img_bytes).decode('utf-8'))
                else:
                    print(f"PDF contains ample selectable text ({len(text)} chars). Skipping image/Vision overhead to avoid model context crowding.")
                
                print(f"PDF processing done: {len(text)} chars, {len(images_base64)} images")

            elif file_name.endswith(".docx"):
                try:
                    print(f"Attempting to extract text from DOCX: {file_name}")
                    text = extract_docx(io.BytesIO(file_content))
                    print(f"DOCX extraction success: {len(text)} chars")
                except Exception as docx_err:
                    print(f"DOCX extraction failed: {str(docx_err)}")
                    text = ""

            elif file_name.endswith(".doc"):
                # .doc files are binary OLE files. python-docx cannot read them.
                # We do a basic string extraction of printable characters as a fallback.
                print(f"Attempting legacy DOC extraction: {file_name}")
                try:
                    # Extract printable characters only to avoid messing up AI/Terminal
                    raw_text = file_content.decode('latin-1', errors='ignore')
                    text = "".join([c for c in raw_text if c.isprintable() or c in "\n\r\t"])
                    print(f"Legacy DOC extraction success: {len(text)} chars")
                except:
                    text = ""

            elif file_name.endswith((".jpg", ".jpeg", ".png")):
                print(f"Optimizing image for Vision upload: {file_name}")
                # Re-encode image to JPEG 80 if it's too large or not JPEG
                img = PILImage.open(io.BytesIO(file_content)).convert('RGB')
                buffer = io.BytesIO()
                img.save(buffer, format="JPEG", quality=80)
                images_base64.append(base64.b64encode(buffer.getvalue()).decode('utf-8'))

            elif file_name.endswith(".txt"):
                text = file_content.decode('utf-8', errors='ignore')
                print(f"TXT extraction success: {len(text)} chars")

            else:
                return Response({"error": "Unsupported file type. Supported: PDF, DOCX, DOC, JPG, JPEG, PNG, TXT"}, status=400)

        except Exception as e:
            import traceback
            print(f"File processing error: {traceback.format_exc()}")
            return Response({"error": f"Error processing file: {str(e)}"}, status=400)

        if not text and not images_base64:
            return Response({"error": "Could not extract any content from the document."}, status=400)

        # print(f"Extracted Text Preview: {text[:500]}...") # Removed to avoid UnicodeEncodeError on Windows terminal
        print(f"Sending {len(text)} chars to {settings.AZURE_OPENAI_DEPLOYMENT}...")

        prompt = f"""Extract all student profile data from this document (transcript or resume).
Return ONLY a valid JSON object with the exact structure below. 

STRICT RULES:
1. CRITICAL EXTRACTION DIRECTIVE: Do NOT skip, omit, summarize, or truncate ANY items. You MUST extract every single educational experience, school term, grade/score, subject, professional internship, full-time job, project, extra-curricular activity, certification, course, and award in its entirety. Full exhaustiveness is mandatory.
2. DO NOT HALLUCINATE. If a piece of information is not present, return an empty string "" or an empty array [].
2. PERSONAL DETAILS: 
   - countryCode: Extract the numeric country code prefix (e.g. "+91" or "91").
   - phoneNumber: Extract ONLY the mobile number without the country code.
   - dob: Look for "DOB" or "Date of Birth". (e.g. "22 March 2003" -> "2003-03-22").
   - gender: "Male", "Female", or "Other".
   - citizenShip: Extract nationality/citizenship (e.g. "Indian").
   - addressline: Extract the street address/house number.
   - city: Extract the city name.
   - zipcode: Extract the 6-digit pin code or zip code.
3. INSTITUTION NAME: Look for the specific college name. 
4. DATES: Use YYYY-MM-DD format. If a date says "Present" or "Current", use "2026-05-09".
5. ACADEMIC LEVEL: Use exactly one of: "High School (8th–12th grade)", "College/Undergraduate", "Postgraduate", "Working Professional". 
   - CRITICAL: If the document contains any Full-time professional work experience (excluding internships), the ACADEMIC LEVEL must be "Working Professional".
   - If the person has completed a Bachelor's degree and is currently in a Master's or has professional experience, it must NOT be "College/Undergraduate".
6. EDUCATIONAL RECORDS: Extract ALL educational experiences (High School, Bachelor, Master, etc.) into the "educational" array.
7. PROFESSIONAL EXPERIENCES: Extract all work experiences into the "professional.experiences" array.
8. HIGH SCHOOL DATA: For each High School record:
   - Set `academicLevel` to exactly "High School (8th–12th grade)".
   - Set `gradeLevel` to the NUMERIC grade only (e.g. 12 for "12th grade" or "Grade 12", 10 for "10th grade"). Do NOT include any text suffix.
   - Extract `board` (e.g. CBSE, ICSE, ISC, IB).
   - Extract `overallPercentage` as the overall/aggregate percentage or score (a plain number, NO % symbol). If per-subject scores are given and no aggregate is stated, leave it empty.
   - Extract all `subjects` with their scores.
9. SUBJECTS: For each subject, extract the name, level (if any), marks obtained (yourTotalScore as a plain number, NO % symbol), and maximum possible marks (highestTotalScore, also a plain number e.g. 100).
10. SCORES & PERCENTILES: Always return numeric scores and percentiles as plain integers or numbers WITHOUT any suffix, symbol, or unit. Do NOT include "%", "th", "rd", "st", "CGPA", "GPA", "/100", etc. Just the number (e.g. "97th Percentile" -> 97, "95%" -> 95, "7.8 CGPA" -> 7.8).
11. PROFESSIONAL EXPERIENCES: 
   - experienceType: Use "Internship", "Full-time", or "Project".
   - currentEmployer: Company name.
   - city: City of the workplace.
   - responsibilities: Bullet points or summary of what was done.
   - achievements: Key results or metrics achieved.
   - startDate/endDate: Extract the month and year (e.g. "May 2023" -> "2023-05-01"). If endDate is "Present", use "2026-05-09".
12. COURSES & CERTIFICATIONS: Extract ALL academic or professional certifications (e.g. Coursera, AWS, NPTEL, Online Courses). Map these to the "courses" array.
13. AWARDS & SCHOLARSHIPS: Extract ALL academic honors, dean's list, scholarships, competition wins (e.g. GSoC, Hackathons, Rank in exams). Map these to the "awards" array.
    - IMPORTANT: Extract these regardless of whether they belong to High School, UG, or PG. They should all go into these global arrays.
    - Note: These items are often found in sections like "ACHIEVEMENTS", "AWARDS", "FELLOWSHIPS", or "CERTIFICATIONS".
14. TEST SCORES: Extract ALL standardized test scores (SAT, ACT, GRE, GMAT, TOEFL, IELTS, Executive Assessment, AP, etc.).
    - ALWAYS extract percentiles if given! Even if they are inside yellow boxes or tables (e.g. "97th Percentile", "96th", "93rd"). Clean them to plain numbers (e.g. 97, 96, 93).
    - For GMAT / GMAT Focus: Extract totalScore, yourPercentile (overall percentile e.g. 96), dataInsightsScore, dataInsightsPercentile (e.g. 90), verbalReasoningScore, verbalReasoningPercentile (e.g. 93), quantitativeReasoningScore, quantitativeReasoningPercentile (e.g. 97).
    - For GRE: Extract totalScore, analyticalWritingScore, analyticalWritingPercentile (e.g. 93), verbalReasoningScore, verbalReasoningPercentile (e.g. 91), quantitativeReasoningScore, quantitativeReasoningPercentile (e.g. 96).
    - For SAT: Extract totalScore, mathYourScore, mathYourPercentile, criticalReadingYourScore (which represents the combined Evidence-Based Reading and Writing score), criticalReadingYourPercentile.
    - For ACT: Extract totalScore, englishYourScore, englishYourPercentile, mathYourScore, mathYourPercentile, readingYourScore, readingYourPercentile, scienceYourScore, scienceYourPercentile.
    - For Executive Assessment: Extract totalScore, integratedReasoningScore, integratedReasoningPercentile, verbalReasoningScore, verbalReasoningPercentile, quantitativeReasoningScore, quantitativeReasoningPercentile.
    - For TOEFL & IELTS: Extract totalScore. Also extract subject-wise scores if available: readingYourScore, listeningYourScore, speakingYourScore, writingYourScore.
    - For AP (Advanced Placement): Extract as testType "AP". Extract the subject name into `testType` (e.g. "AP Physics") and the score (1-5) into `yourScore`.
    - For Others: Extract yourScore.
15. SKILLS & PROJECTS: Extract skills and major projects. Map these to "additional" section if no specific section exists.

JSON STRUCTURE:
{{
  "personalDetails": {{ "firstName": "", "lastName": "", "email": "", "countryCode": "", "phoneNumber": "", "dob": "YYYY-MM-DD", "gender": "", "citizenShip": "", "addressline": "", "city": "", "zipcode": "", "mothersProfession": "", "fathersProfession": "" }},
  "educational": [
    {{
       "academicLevel": "High School (8th\u201312th grade)", "gradeLevel": 12, "institutionName": "", "city": "", "yearOfCompletion": "YYYY-MM-DD",
       "overallPercentage": "", "maximumPossibleGPA": "",
       "degree": "", "major": "", "startYear": "", "endYear": "",
       "board": "",
       "subjects": [
          {{ "subject": "", "level": "", "yourTotalScore": "", "highestTotalScore": "100" }}
       ]
    }}
  ],
  "courses": [
     {{ "courseType": "", "description": "", "year": "", "awards": "", "duration": "", "location": "" }}
  ],
  "awards": [
     {{ "nameOfHonorReceived": "", "description": "", "levelOfCompetitiveness": "", "numberOfParticipants": "", "year": "" }}
  ],
  "testScores": [ 
    {{ 
      "testType": "", "testDate": "YYYY-MM-DD", "totalScore": "", "yourScore": "", "yourPercentile": "",
      "mathYourScore": "", "mathYourPercentile": "", "criticalReadingYourScore": "", "criticalReadingYourPercentile": "",
      "analyticalWritingScore": "", "analyticalWritingPercentile": "", "verbalReasoningScore": "", "verbalReasoningPercentile": "", "quantitativeReasoningScore": "", "quantitativeReasoningPercentile": "",
      "dataInsightsScore": "", "dataInsightsPercentile": "", "englishYourScore": "", "englishYourPercentile": "", "readingYourScore": "", "readingYourPercentile": "", "scienceYourScore": "", "scienceYourPercentile": "",
      "integratedReasoningScore": "", "integratedReasoningPercentile": "", "listeningYourScore": "", "listeningYourPercentile": "", "speakingYourScore": "", "speakingYourPercentile": "", "writingYourScore": "", "writingYourPercentile": ""
    }} 
  ],
  "professional": {{
     "experiences": [ 
        {{ "experienceType": "", "currentEmployer": "", "startDate": "YYYY-MM-DD", "endDate": "YYYY-MM-DD", "jobTitle": "", "city": "", "industrySector": "", "responsibilities": "", "achievements": "" }} 
     ]
  }},
  "extraCurricular": [
     {{ "activityType": "", "startDate": "YYYY-MM-DD", "endDate": "YYYY-MM-DD", "positionHeld": "", "awardsCertifications": "", "description": "" }}
  ],
  "skills": [],
  "projects": [ {{ "title": "", "description": "", "startDate": "", "endDate": "" }} ]
}}

Document content:
{text[:100000]}
"""

        try:
            messages = [
                {"role": "system", "content": "You are an expert academic data extractor. Extract data from the provided text and images, and return valid JSON only."}
            ]
            
            if images_base64:
                user_content = [
                    {"type": "text", "text": prompt},
                ]
                for img in images_base64:
                    user_content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{img}", "detail": "high"}
                    })
                messages.append({"role": "user", "content": user_content})
            else:
                messages.append({"role": "user", "content": prompt})

            response = client.chat.completions.create(
                model=settings.AZURE_OPENAI_DEPLOYMENT,
                response_format={"type": "json_object"},
                messages=messages,
                temperature=0
            )
            print("Azure OpenAI response received.")
            content = response.choices[0].message.content
            # print(f"--- AI RESPONSE START ---\n{content}\n--- AI RESPONSE END ---")
            
            # Save raw AI response to a local file for debugging
            try:
                scratch_dir = os.path.join(settings.BASE_DIR, 'scratch')
                os.makedirs(scratch_dir, exist_ok=True)
                with open(os.path.join(scratch_dir, 'last_ai_response.json'), 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"Saved raw AI response to local scratch file.")
            except Exception as save_err:
                print(f"Failed to save AI response to scratch: {save_err}")
                
            parsed = json.loads(content)
            print(f"Successfully parsed AI response with keys: {list(parsed.keys())}")
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"Azure OpenAI error: {error_trace}")
            return Response({
                "error": f"AI processing failed: {str(e)}",
                "trace": error_trace if settings.DEBUG else None
            }, status=500)

        return Response(parsed)